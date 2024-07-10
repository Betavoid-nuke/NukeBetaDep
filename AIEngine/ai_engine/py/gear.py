# (c) CID Research LLC in partnership with Betavoid

import os
import sys
import math
import random
import datetime
import numpy as np
from pathlib import Path
from database import data
from ai_engine.py.addon import Input
import misc.cad.script as codegen
from config import items
try:
  from docx import Document
except:
  import subprocess
  subprocess.run(["pip", "install", "python-docx"], check=True)
  from docx import Document

try:
  from tabulate import tabulate
except:
  import subprocess
  subprocess.run(["pip", "install", "tabulate"], check=True)
  from tabulate import tabulate

# Gear Designer
class GearDesigner(object):
    def __init__(self, *args, **kwargs):
        self.args = args
        self.kwargs = kwargs
        self.varfile = None
        self.db = data.Database()
        self.db.ACRONYM = items.ACRONYM

    def compute_circle_diameter(self):
        self.data.cd = (0.8 * sum(self.data.volume[:2])//2)

    def compute_diametric_pitch(self):
        pd = self.data.pd or self.data.volume[0]
        self.data.dp = self.data.nT/pd

    def compute_pitch_diameter(self):
        self.data.pd = (self.data.nT - 1.8)/self.data.dp

    def compute_shear_stress(self):
        if math.pow(self.data.pd, 3) == 0:
            self.data.pd = self.data.volume[0]+random.random()
        self.data.ss = (16 * self.data.torque * self.data.ka)/(math.pi * math.pow(self.data.pd, 3) * self.data.kf)

    def compute_force_of_safety(self):
        self.data.fos = self.data.ms/self.data.ss

    def compute_length_of_splines(self):
        self.data.los = self.data.d * 0.2

    def adjust_value(self, low, high):
        to_search = high if self.data.fos < self.db.fos_range[0] else low
        if len(to_search) == 0:
            to_search = low + high
        value, low, high = self.adjust(to_search)
        return value, low, high

    def adjust(self, values):
        m = len(values)//2
        x0 = values[:m]
        x1 = values[m:]
        return values[m], x0, x1

    def mm2inch(self, x):
        return x/25.4

    def convert_units(self):
        self.data.torque *= 8.850745767378
        self.data.volume = [self.mm2inch(x) for x in self.data.volume]

    def assumption(self, min=0, max=100, mode=30):
        x = np.arange(min,mode)
        x0 = [i for i in x if i<mode]
        x1 = [i for i in x if i>mode]
        return mode, x0, x1

    def get_likelihood(self, query, values):
        i,j = np.unravel_index(np.argmax(query), query.shape)
        x = values[i+1,j+1]
        query[i,j] = 0.0
        return x

    def get_decider(self):
        product = {
            "jet": 2.1,
            "machine": 1.1,
            "car": 0.3,
            "bike": 0.2,
            "bicycle": 0.1, }
        price = (self.data.budget/1000)
        for k,v in product.items():
          if k in self.data.of:
            break
        scale = int(np.round(price + v, 0))
        return scale

    def query(self):
        picker = self.get_decider()
        self.data.ms = self.db.mechanical_property_dict['tensile_strength']['yield'] # allowable stress
        self.data.ka = self.db.ka_range[picker]
        self.data.kf = self.db.kf_range[picker]
        self.data.km = self.db.km_range[picker]
        self.data.kw = self.db.kw_range[picker]
        self.db.fos_range = [(picker+1)*2, (picker+1.7)*2]
        return picker

    def measure(self):
        self.compute_diametric_pitch()
        self.compute_pitch_diameter()
        self.compute_shear_stress()
        self.compute_force_of_safety()
        return

    def compute_diamters(self):
      md1 = (self.data.nT - 1.8)/self.data.dp
      return dict(
          MD0 = (self.data.nT + 1.8)/self.data.dp, # major diameter, internal
          MD1 = (self.data.nT + 1)/self.data.dp, # major diameter, external
          md0 = (self.data.nT- 1)/self.data.dp, # minor diameter, internal
          md1 = (self.data.nT - 1.8)/self.data.dp, # minor diameter, external - thru + pitch
          md2 = (self.data.nT - 2)/self.data.dp, # minor diameter, external - pitch + fine
          Add = 1/self.data.dp, # Addendum
          Ddd = 1.25/self.data.dp, # Dedendum
          Ddds = 1.35/self.data.dp, # Dedendum (shaved or ground teeth)
          wkd = 2/self.data.dp, # working depth
          whd = 2.25/self.data.dp, # whole depth (preffered)
          whds = 2.35/self.data.dp, # whole depth (shaved or ground teeth)
          clr = 0.25/self.data.dp, # Clearance
          clrs = 0.35/self.data.dp, # Clearance (shaved or ground teeth)
          flr = 0.3/self.data.dp, # filet radius
          pd = self.data.nT/self.data.dp, #? repeated but different formular and values, WHY FORMULAR CHANGE?
          psd = (self.data.nT+2)/self.data.dp, # piutside diameter
          rtd = (self.data.nT-2.5)/self.data.dp, # root diameter (preffered)
          rtds = (self.data.nT-2.7)/self.data.dp, # root diameter (shaved or ground teeth)
          ctb = 1.5708/self.data.dp, # circle thickness basic
          cc = md1 + 1.2,
          cm= ((1-(79.9-md1)/79.9)/2)+1,
          lm = 1-(79.9-md1)/79.9,
          D = (self.data.pd/self.data.nT) * 2.5,
        )

    def process(self, inputs:dict):
        self.__init__()
        # Select optimal input template (ML)
        self.data = Input.get_optimal_template(inputs)
        print("PROCESSED")
        self.data.save = self.save
        # Update input template based on prompts (Maths)
        for k,v in inputs.items():
            self.data[k] = v
        # Preprocess inputs (ML)
        self.convert_units()
        # fix basic assumptions
        picker = self.query()
        pa, pa_low, pa_high = self.assumption(min=self.db.pa_range[0], max=self.db.pa_range[2], mode=self.db.pa_range[1]) # Pressure angle
        nT, teeth_low, teeth_high = self.assumption(min=self.db.teeth_range[0], max=self.db.teeth_range[2], mode=self.db.teeth_range[1]) # No of Teeth
        self.data.pa = pa
        self.data.nT = nT
        # initial calculations
        self.compute_circle_diameter()
        m = self.measure()
        #5. Optimization loop (DL)
        while not(self.data.fos >= self.db.fos_range[0] and self.data.fos <= self.db.fos_range[1]):
            idx = random.randint(0, 3)
            if idx == 0:
                nT, teeth_low, teeth_high = self.adjust_value(teeth_low, teeth_high)
                self.data.nT = nT
            elif idx == 1:
                self.data.ka = self.db.ka_range[picker-1]
            elif idx == 2:
                self.data.kf = self.db.kf_range[picker-1]
            elif idx == 4:
                pa, pa_low, pa_high = self.adjust_value(pa_low, pa_high)
                self.data.pa = pa
            m = self.measure()
        self.data.achieved = True
        out = self.compute_diamters()
        for k,v in out.items():
            self.data[k] = v
        # coding
        cdg = codegen.CodeGen()
        cdg.code(self.data)
        self.data["code"] = self.code = cdg
        print("Completed ", datetime.datetime.now(), ". \n >> You can query your model for any VAR using 'model.data.VAR'. \n >> Use 'model.db.ACRONYM' to see the VAR name.")
        return self

    def save(self, filepath=None):
        if not (filepath or self.varfile):
            sys.exit("Error: No variable filepath provided!, ...give the filepath to save the output variables")
        if filepath:
            self.varfile = Path(filepath).with_suffix(".docx")
        os.makedirs(os.path.dirname(self.varfile), exist_ok=True)
        output = self.present_vars(self.data)
        doc = Document()
        doc.add_heading('Tabulated Variable Data', level=1)
        # Add a table with the same shape as the NumPy array
        rows, cols = output.shape
        table = doc.add_table(rows=rows, cols=cols)
        # Populate the table with NumPy array values
        for i in range(rows):
            for j in range(cols):
                cell = table.cell(i, j)
                cell.text = str(output[i, j])
        # doc.add_paragraph(output)
        doc.save(self.varfile)
        print("Output variables saved @ ", self.varfile)

    def __str__(self, data=None):
        if data==None: data = self.data
        output = self.tabulate_vars(data)
        return output

    def tabulate_vars(self, data):
        data = self.present_vars(data)
        table = (tabulate(data[1:], headers=data[0], tablefmt='grid'))
        return table

    def present_vars(self, data):
        head = [['ACRONYM', 'ATTRIBUTE', 'VALUE']]
        body = [(k, self.db.ACRONYM[k], v) for k, v in sorted(data.items()) if k in self.db.ACRONYM.keys()]
        head.extend(body)
        return np.array(head)

